"""DOCX ingestor — paragraph-level extraction via python-docx."""

from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator

from doc2dataset.core.registry import register_ingestor
from doc2dataset.exceptions import CorruptedFileError
from doc2dataset.ingestion.abstract_ingestor import AbstractIngestor
from doc2dataset.processing.metadata import DocumentFragment, FragmentMetadata

logger = logging.getLogger(__name__)


@register_ingestor(".docx")
class DocxIngestor(AbstractIngestor):
    """Extract text from ``.docx`` files one paragraph at a time.

    Tables are serialised as pipe-delimited rows so that tabular data is
    not silently discarded.
    """

    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    def ingest(self, path: Path) -> Iterator[DocumentFragment]:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required — pip install python-docx"
            ) from exc

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise CorruptedFileError(f"Cannot open DOCX {path}: {exc}") from exc

        # ── Body paragraphs ────────────────────────────────────────────────
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            yield DocumentFragment(
                text=text,
                metadata=FragmentMetadata(
                    source_path=str(path),
                    page_or_section=idx,
                    timestamp=datetime.now(timezone.utc).isoformat(),
                    extra={"style": para.style.name if para.style else ""},
                ),
            )

        # ── Tables → pipe-delimited text ───────────────────────────────────
        for tbl_idx, table in enumerate(doc.tables):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                yield DocumentFragment(
                    text="\n".join(rows),
                    metadata=FragmentMetadata(
                        source_path=str(path),
                        page_or_section=f"table-{tbl_idx}",
                        timestamp=datetime.now(timezone.utc).isoformat(),
                    ),
                )
